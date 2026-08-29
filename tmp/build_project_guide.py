from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path

ROOT = Path(r"C:\Users\facun\Desktop\QUIMICA - V3")
OUT = ROOT / "Guia_Arquitectura_Quimicas_King.docx"

BLUE = "1F4E79"
LIGHT = "DCE6F1"
PALE = "F4F7FA"
RED = "9B1C1C"
GREEN = "1F5E3B"

def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    tcPr.append(shd)

def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in('w:tcMar')
    if tcMar is None:
        tcMar = OxmlElement('w:tcMar')
        tcPr.append(tcMar)
    for m, v in [('top', top), ('start', start), ('bottom', bottom), ('end', end)]:
        node = tcMar.find(qn(f'w:{m}'))
        if node is None:
            node = OxmlElement(f'w:{m}')
            tcMar.append(node)
        node.set(qn('w:w'), str(v)); node.set(qn('w:type'), 'dxa')

def set_width(cell, width):
    tcPr = cell._tc.get_or_add_tcPr()
    tcW = tcPr.find(qn('w:tcW'))
    if tcW is None:
        tcW = OxmlElement('w:tcW'); tcPr.append(tcW)
    tcW.set(qn('w:w'), str(width)); tcW.set(qn('w:type'), 'dxa')

def style_table(table, widths, header=True):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    for row_i, row in enumerate(table.rows):
        for i, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            set_width(cell, widths[i])
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(2)
                for run in p.runs:
                    run.font.size = Pt(9.5)
            if header and row_i == 0:
                shade(cell, LIGHT)
                for run in cell.paragraphs[0].runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor.from_string(BLUE)

def para(doc, text='', bold_prefix=None, color=None, size=None, style=None):
    p = doc.add_paragraph(style=style)
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix); r.bold = True
        p.add_run(text[len(bold_prefix):])
    else:
        p.add_run(text)
    if color or size:
        for r in p.runs:
            if color: r.font.color.rgb = RGBColor.from_string(color)
            if size: r.font.size = Pt(size)
    return p

def bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(item)

def add_table(doc, headers, rows, widths=(2600, 6760)):
    table = doc.add_table(rows=1, cols=len(headers))
    for i, h in enumerate(headers): table.rows[0].cells[i].text = h
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row): cells[i].text = value
    style_table(table, widths)
    doc.add_paragraph()
    return table

doc = Document()
sec = doc.sections[0]
sec.top_margin = Inches(0.75); sec.bottom_margin = Inches(0.7)
sec.left_margin = Inches(0.75); sec.right_margin = Inches(0.75)

styles = doc.styles
styles['Normal'].font.name = 'Calibri'; styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
styles['Normal'].font.size = Pt(10.5)
styles['Normal'].paragraph_format.space_after = Pt(6)
styles['Normal'].paragraph_format.line_spacing = 1.15
for name, size, color in [('Title', 25, BLUE), ('Heading 1', 16, BLUE), ('Heading 2', 13, BLUE), ('Heading 3', 11.5, '1F4E79')]:
    s = styles[name]; s.font.name = 'Calibri'; s._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
    s.font.size = Pt(size); s.font.color.rgb = RGBColor.from_string(color)

# Cover
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('GUÍA MAESTRA DEL PROYECTO'); r.bold = True; r.font.size = Pt(25); r.font.color.rgb = RGBColor.from_string(BLUE)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Químicas King - mapa de archivos, lógica y reglas para trabajar con IA'); r.font.size = Pt(14); r.font.color.rgb = RGBColor.from_string('4A5568')
doc.add_paragraph()
table = doc.add_table(rows=1, cols=1); table.cell(0,0).text = 'PROPÓSITO\nExplicar la web desde el archivo HTML de inicio hasta cada pantalla y componente, de forma que una persona no técnica pueda pedir cambios con seguridad.'
style_table(table, (9360,), header=False); shade(table.cell(0,0), PALE)
doc.add_paragraph()
para(doc, 'Cómo usar esta guía', style='Heading 1')
bullets(doc, [
    'Primero ubicá la necesidad en la sección “Dónde cambiar cada cosa”.',
    'Después leé las dependencias: un mismo dato puede alimentar catálogo, ficha, carrito y WhatsApp.',
    'Al usar una IA, pegá el prompt de la última sección y pedile que modifique solamente los archivos indicados.',
    'No edites manualmente src/data/products.js: es un archivo generado desde Excel.'
])
para(doc, 'Estado del relevamiento: estructura analizada el 29 de agosto de 2026. Esta guía describe el código actualmente presente, incluidos tres puntos de mantenimiento detectados.', color='4A5568', size=9)
doc.add_page_break()

doc.add_heading('1. Vista general: qué tipo de proyecto es', level=1)
para(doc, 'Es una tienda/catálogo de productos de limpieza creada con React y Vite. No hay backend ni base de datos: los productos vienen de un Excel y se convierten a un archivo JavaScript. El pedido termina en WhatsApp; el carrito se guarda localmente en el navegador del cliente.')
add_table(doc, ['Capa', 'Responsabilidad'], [
    ('Base web', 'index.html contiene el único contenedor HTML (#root) e inicia React.'),
    ('Inicio técnico', 'src/main.jsx monta la aplicación, carga estilos globales y habilita las rutas.'),
    ('Reglas globales', 'src/App.jsx activa el carrito para toda la web; AppRoutes.jsx decide qué página mostrar.'),
    ('Pantallas', 'src/pages contiene Home, Catálogo, Ficha de producto, Promociones y Nosotros.'),
    ('Piezas visuales', 'src/components contiene secciones reutilizables: hero, cards, navbar, carrito, etc.'),
    ('Datos comerciales', 'Excel -> generateProducts.cjs -> src/data/products.js -> catálogo, ficha, carrito y WhatsApp.'),
    ('Estilos e imágenes', 'CSS en src/styles y junto a los componentes; imágenes públicas en public/.')
])
doc.add_heading('2. Mapa de arranque: desde HTML hasta Home', level=1)
para(doc, 'Seguí este recorrido cuando quieras entender de dónde “nace” la pantalla inicial:')
flow = [
    'index.html  →  incluye <div id="root"> y carga /src/main.jsx',
    'src/main.jsx  →  carga CSS global + BrowserRouter + <App />',
    'src/App.jsx  →  envuelve toda la web con CartProvider y carga AppRoutes',
    'src/routes/AppRoutes.jsx  →  para la ruta / muestra <Home />',
    'src/pages/Home.jsx  →  compone las secciones visuales del inicio',
    'components/home/*.jsx  →  dibujan cada bloque del Home'
]
for step in flow:
    p = doc.add_paragraph(style='List Number'); p.add_run(step)
para(doc, 'Importante: BrowserRouter permite cambiar de página sin recargar. Las URL relevantes son /, /catalogo, /promociones, /nosotros y /producto/:slug.', bold_prefix='Importante:')

doc.add_heading('3. Rutas y pantallas', level=1)
add_table(doc, ['URL', 'Archivo principal', 'Qué muestra / de qué depende'], [
    ('/', 'src/pages/Home.jsx', 'Inicio: hero, categorías, productos destacados, estadísticas, proceso de compra, reseñas, pie y carrito.'),
    ('/catalogo', 'src/pages/Catalog.jsx', 'Listado de familias de productos, búsqueda, categorías, subcategorías y carrito.'),
    ('/promociones', 'src/pages/Promotions.jsx', 'Reutiliza Catalog.jsx pero le pasa únicamente productos con featured = true.'),
    ('/producto/:slug', 'src/pages/Product.jsx', 'Ficha de una variante específica; selector de presentaciones, cantidad, carrito, WhatsApp y relacionados.'),
    ('/nosotros', 'src/pages/AboutPage.jsx', 'Información institucional, preguntas, envíos, política, contacto y pie.'),
    ('cualquier otra', 'src/routes/AppRoutes.jsx', 'Redirige al inicio para evitar una página inexistente.')
], (1700, 2700, 4960))
doc.add_page_break()

doc.add_heading('4. Home: piezas, orden y lugar correcto para editar', level=1)
para(doc, 'Home.jsx es el “director de orquesta”: define el orden de las secciones y abre/cierra el carrito. Cada bloque tiene su propio componente para que se pueda editar una sección sin alterar las demás.')
add_table(doc, ['Orden visible', 'Componente', 'Editar aquí cuando quieras cambiar...'], [
    ('1. Cabecera visual', 'components/home/Hero.jsx + Hero.css', 'imagen de portada, título, llamados a acción o navegación superior.'),
    ('2. Categorías', 'components/home/Categorysection.jsx + CSS', 'las tres categorías destacadas y el enlace que lleva al catálogo.'),
    ('3. Destacados', 'components/home/FeaturedProducts.jsx y FeaturedCard.jsx', 'carrusel/animación de productos destacados y su tarjeta.'),
    ('4. Estadísticas', 'components/home/StatsStrip.jsx + CSS', 'contadores; Home calcula productos activos, categorías y años.'),
    ('5. Proceso', 'components/home/DeliveryProcess.jsx + CSS', 'pasos para comprar o textos de envío.'),
    ('6. Reseñas y mapa', 'ReviewsSection.jsx, LocationsMap.jsx, data/reviews.js, data/locations.js', 'opiniones y sucursales/direcciones del minimapa.'),
    ('7. Pie', 'components/home/Footer.jsx + CSS', 'datos de contacto, enlaces, redes y crédito final.'),
    ('Transversal', 'common/FadeInSection.jsx', 'animación de entrada; no hace falta tocarlo para cambiar contenido.')
], (1450, 3500, 4410))
doc.add_heading('5. El núcleo comercial: Excel, grupos y variantes', level=1)
para(doc, 'La fuente de verdad comercial es el Excel público. generateProducts.cjs lo lee y reconstruye src/data/products.js. Esa separación evita escribir productos a mano, pero obliga a respetar el flujo de generación.')
add_table(doc, ['Elemento', 'Qué significa', 'Regla que no se debe romper'], [
    ('GRUPO', 'Identifica una familia: por ejemplo, el mismo producto en 1 L y 5 L.', 'Todas las variantes de una familia deben tener exactamente el mismo GRUPO.'),
    ('VARIANTE', 'Texto visible de la presentación, por ejemplo “1 litro”.', 'Debe ser comprensible para el cliente y estar presente en cada fila.'),
    ('SLUG', 'Identificador único de la URL, del carrito y de la ficha.', 'No se repite nunca. El generador detiene el proceso si encuentra duplicados.'),
    ('NOMBRE FOTO', 'Archivo de imagen dentro de public/products/.', 'El nombre y extensión deben coincidir exactamente con la imagen subida.'),
    ('salePrice', 'Precio numérico que se usa para cálculos.', 'Siempre usarlo para subtotal y total.'),
    ('price', 'Precio formateado para mostrar, por ejemplo $6.400.', 'Solo para texto visual; no usarlo para multiplicar.')
], (1850, 2900, 4610))
para(doc, 'Flujo de datos: public/Plantilla_Quimica_Bethel_2.0.xlsx → generateProducts.cjs → src/data/products.js → products (familias) y productVariants (variantes individuales).', bold_prefix='Flujo de datos:')
bullets(doc, [
    'products se usa en Home, Catalog y Promociones. Cada familia aparece una vez y su precio es el menor de sus variantes.',
    'productVariants se usa en Product.jsx. Permite abrir una URL exacta y elegir otra presentación dentro de la misma familia.',
    'Después de cambiar Excel, se debe ejecutar el script de generación. Nunca corregir el resultado generado a mano porque se perderá en la próxima generación.'
])

doc.add_heading('6. Catálogo y ficha: la lógica de compra', level=1)
add_table(doc, ['Zona', 'Archivo(s)', 'Lógica esencial'], [
    ('Filtros', 'pages/Catalog.jsx', 'Mantiene estado de búsqueda, categoría y subcategoría. La subcategoría depende de la categoría seleccionada. Se limpia al cambiar categoría.'),
    ('Grilla', 'catalog/ProductGrid.jsx', 'Recibe productos ya filtrados y renderiza una ProductCard por familia.'),
    ('Card', 'catalog/ProductCard.jsx + CSS', 'Muestra “Desde” y cantidad de presentaciones. Navega a la ficha; no compra directamente.'),
    ('Ficha', 'pages/Product.jsx + Product.css', 'Busca una variante por slug, reúne variantes del mismo group, calcula subtotal con salePrice y arma el mensaje de WhatsApp.'),
    ('Cambio de variante', 'pages/Product.jsx', 'navigate(`/producto/${variant.slug}`, { replace: true }) cambia URL sin agregar pasos al botón Volver.'),
    ('Relacionados', 'pages/Product.jsx', 'Filtra misma categoría y group distinto; así no recomienda otra versión del mismo artículo.')
], (1500, 3000, 4860))
doc.add_heading('7. Carrito y WhatsApp', level=1)
para(doc, 'El carrito es global: CartProvider lo activa antes de cargar cualquier página. Se persiste en localStorage con la clave bethel-cart, por lo que el pedido permanece si el visitante actualiza la página en el mismo navegador.')
add_table(doc, ['Archivo', 'Rol'], [
    ('context/CartContext.jsx', 'Fuente única del carrito: agrega, borra, sube/baja cantidades, calcula cantidad total y total monetario. Identifica ítems por slug.'),
    ('components/cart/CartButton.jsx', 'Botón flotante que muestra cuántas unidades hay y abre el panel.'),
    ('components/cart/CartDrawer.jsx', 'Panel lateral: lista ítems, muestra total, permite vaciar y genera el pedido de WhatsApp.'),
    ('components/cart/CartItem.jsx', 'Una línea del carrito: imagen, nombre, variante, precio y controles de cantidad.'),
    ('styles/cart.css', 'Diseño del panel; mantener un solo contenedor .cart-items para evitar doble desplazamiento en celular.')
], (3000, 6360))
para(doc, 'Contrato del carrito: las funciones de CartContext reciben slug. Si se modifica CartItem, los botones deben llamar decreaseQuantity(item.slug), increaseQuantity(item.slug) y removeFromCart(item.slug).', bold_prefix='Contrato del carrito:', color=RED)
doc.add_page_break()

doc.add_heading('8. Inventario de archivos: qué hay en cada carpeta', level=1)
para(doc, 'Esta es la jerarquía útil de Visual Studio Code. Los CSS que se nombran junto a un componente controlan principalmente esa pieza; los CSS de src/styles afectan áreas compartidas.')
add_table(doc, ['Carpeta / archivo', 'Función'], [
    ('index.html', 'Base HTML con #root; no contiene el diseño de la web.'),
    ('package.json', 'Comandos del proyecto y dependencias: React, Vite, router, iconos, animaciones y lectura de Excel.'),
    ('generateProducts.cjs', 'Transforma la planilla de productos en datos para React y valida SLUGs duplicados.'),
    ('public/', 'Recursos públicos: logo, hero, foto institucional, favicon y public/products/ con fotos de artículos.'),
    ('src/main.jsx', 'Puerta técnica de entrada; importa estilos globales, router y App.'),
    ('src/App.jsx', 'Conecta el carrito global con las rutas.'),
    ('src/routes/AppRoutes.jsx', 'Mapa de URLs hacia páginas.'),
    ('src/pages/', 'Pantallas completas: Home, Catalog, Product, Promotions, AboutPage.'),
    ('src/components/home/', 'Secciones visuales del inicio y algunas usadas en Nosotros.'),
    ('src/components/catalog/', 'Buscador, grilla y tarjeta de catálogo.'),
    ('src/components/cart/', 'Botón, panel e ítems del carrito.'),
    ('src/components/about/', 'Hero, envío y política de revisión de Nosotros.'),
    ('src/components/common/', 'Piezas reutilizables: logo, encabezado de sección y animación.'),
    ('src/components/layout/', 'Navbar reutilizable.'),
    ('src/context/', 'Estado compartido del carrito.'),
    ('src/data/', 'Datos: productos generados, reseñas y ubicaciones.'),
    ('src/styles/', 'Estilos generales: global, responsive, catálogo, búsqueda, carrito, bandas y layout.'),
    ('src/utils/productGroups.js', 'Utilidad antigua de agrupación. Actualmente no se importa en el flujo visible.')
], (3000, 6360))
doc.add_heading('9. Componentes de contenido e institucionales', level=1)
add_table(doc, ['Grupo', 'Archivos principales', 'Uso'], [
    ('Comunes', 'Logo.jsx/CSS, SectionHeader.jsx/CSS, FadeInSection.jsx', 'Marca, títulos estandarizados y animación de aparición.'),
    ('Navegación', 'layout/Navbar.jsx/CSS', 'Menú desktop/móvil; Hero y AboutHero lo incorporan.'),
    ('Nosotros', 'AboutHero.jsx/CSS, Deliveryinfo.jsx/CSS, ReviewPolicy.jsx/CSS', 'Encabezado, zonas/costos y política de revisión.'),
    ('Contenido auxiliar', 'AboutUs, FAQ, PolicySection, ContactSection, DeliverySection', 'Secciones de texto del sitio. AboutPage usa AboutUs, FAQ, PolicySection y ContactSection; DeliverySection parece ser una alternativa no montada.'),
    ('Reseñas/mapa', 'ReviewsSection.jsx/CSS, LocationsMap.jsx, data/reviews.js, data/locations.js', 'Carrusel de reseñas y datos de sucursales editables.')
], (2100, 4100, 3160))

doc.add_heading('10. Dónde cambiar cada cosa sin romper nada', level=1)
add_table(doc, ['Quiero cambiar...', 'Editar primero', 'Luego comprobar'], [
    ('Un precio, stock, foto, categoría o variante', 'Excel y la imagen en public/products/', 'Regenerar products.js; abrir catálogo, ficha y carrito.'),
    ('Textos / botones del inicio', 'El componente home específico (Hero, DeliveryProcess, Footer, etc.)', 'Su CSS asociado y móvil.'),
    ('Orden de secciones del Home', 'pages/Home.jsx', 'Que las importaciones sigan existiendo y el carrito continúe al final.'),
    ('Filtros / buscador del catálogo', 'pages/Catalog.jsx', 'Categoría, subcategoría y búsqueda combinadas, también en móvil.'),
    ('Tarjeta de catálogo', 'catalog/ProductCard.jsx + ProductCard.css', 'No agregar compra directa: debe ir a la ficha para elegir variante.'),
    ('Ficha de producto', 'pages/Product.jsx + Product.css', 'URLs por slug, cambio de variantes, subtotal, WhatsApp y relacionados.'),
    ('Reglas de compra / total', 'context/CartContext.jsx y CartDrawer.jsx', 'Usar salePrice numérico y conservar slug como identificador.'),
    ('Una ubicación', 'data/locations.js', 'Que mapQuery sea una dirección o coordenada válidas.'),
    ('Una reseña', 'data/reviews.js', 'id único, stars, author y text.'),
    ('Diseño global o celular', 'styles/global.css, responsive.css y CSS de la sección', 'No corregir un problema local solo con global.css sin revisar efectos en otras páginas.')
], (2450, 3450, 3460))
doc.add_page_break()

doc.add_heading('11. Reglas de seguridad para trabajar con inteligencias artificiales', level=1)
bullets(doc, [
    'Pedir primero que la IA lea los archivos involucrados y explique qué tocará. No permitir cambios “a ciegas”.',
    'Pedir cambios mínimos: una necesidad visual no debería reescribir el router, el carrito ni la generación de productos.',
    'No alterar los nombres/contratos de product.slug, product.group, product.variantLabel, product.salePrice, product.price o product.image sin actualizar todos los consumidores.',
    'Usar salePrice para operar matemáticamente; price solo para mostrar. No reintroducir cost, margin ni unitPrice.',
    'No editar src/data/products.js directamente: se regenera y se perdería el cambio.',
    'Antes de finalizar, solicitar una comprobación de navegación, cambio de variante, carrito, WhatsApp y diseño móvil.',
    'Pedir una lista final de archivos modificados y el motivo de cada cambio.'
])
doc.add_heading('12. Alertas de mantenimiento detectadas', level=1)
para(doc, 'Estas observaciones describen el estado actual. No se aplicó ninguna modificación; se incluyen para evitar que una IA consolide o esconda el problema.', color=RED)
add_table(doc, ['Prioridad', 'Archivo', 'Observación / acción recomendada'], [
    ('Alta', 'components/cart/CartItem.jsx', 'CartContext trabaja con slug, pero CartItem llama funciones con item.id. Los controles de cantidad y eliminar deberían usar item.slug; hoy podrían fallar o afectar incorrectamente ítems agrupados.'),
    ('Alta', 'components/home/FeaturedCard.jsx', 'Calcula el subtotal con cost y margin, campos que ya no existen en el formato actual. Debe pasar a Number(product.salePrice) * quantity y mantener la lógica de variantes.'),
    ('Media', 'utils/productGroups.js', 'Usa unitPrice y no forma parte del flujo actual. Está desactualizado; revisar antes de reutilizarlo o eliminarlo tras confirmar que no se necesita.'),
    ('Baja', 'ProductGrid.jsx / llamadas', 'Recibe openCart, pero no lo usa. Es una propiedad sobrante que se puede limpiar en una tarea separada.'),
    ('Baja', 'TopNavigation.css y DeliverySection.jsx', 'Existen archivos que no parecen montados en las páginas actuales. No borrarlos sin una búsqueda completa de importaciones y revisión visual.')
], (1000, 2850, 5510))

doc.add_heading('13. Prompt listo para copiar y pegar a una IA', level=1)
prompt = '''Actuá como desarrollador cuidadoso de este proyecto React/Vite llamado Químicas King. Antes de modificar nada, leé los archivos vinculados a la tarea y explicame en lenguaje simple: (1) qué archivo controla esa parte visual, (2) qué datos o componentes dependen de él y (3) qué archivos vas a cambiar.

Arquitectura obligatoria: index.html carga src/main.jsx; main.jsx monta App con BrowserRouter; App.jsx envuelve AppRoutes con CartProvider; AppRoutes dirige a Home, Catalog, Promotions, Product y AboutPage. El Home se arma en src/pages/Home.jsx con componentes de src/components/home. El catálogo usa familias de productos desde src/data/products.js. La ficha usa productVariants por slug.

Reglas comerciales no negociables:
- La fuente de productos es public/Plantilla_Quimica_Bethel_2.0.xlsx. generateProducts.cjs genera src/data/products.js. No edites products.js a mano.
- GRUPO identifica una familia de variantes; VARIANTE es la presentación visible; SLUG es único y se usa en URL y carrito.
- products contiene familias agrupadas; productVariants contiene cada variante individual.
- Para cálculos usar Number(product.salePrice). price es solo texto visual. No usar cost, margin ni unitPrice.
- El carrito identifica artículos por slug. Una variante de 1 L y otra de 5 L son ítems distintos.
- La ProductCard del catálogo solo debe llevar a la ficha: primero se elige la variante y después se agrega al carrito.
- Al cambiar variante en Product.jsx, conservar navegación por /producto/:slug con replace: true.

Forma de trabajo: aplicá el cambio mínimo posible, no reestructures archivos no relacionados, conservá la experiencia móvil, no cambies nombres de datos compartidos sin actualizar todos sus usos. Al terminar, indicá los archivos modificados, explicá por qué, y verificá catálogo, ficha, cambio de variante, carrito, total, WhatsApp y móvil. Si detectás una inconsistencia existente, separala de la tarea y pedime autorización antes de corregirla.'''
box = doc.add_table(rows=1, cols=1); box.cell(0,0).text = prompt; style_table(box, (9360,), header=False); shade(box.cell(0,0), PALE)
for p in box.cell(0,0).paragraphs:
    for r in p.runs: r.font.name='Consolas'; r.font.size=Pt(8.5)
para(doc, 'Fin de la guía.', color='4A5568', size=9)

doc.save(OUT)
print(OUT)
